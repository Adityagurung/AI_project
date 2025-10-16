"""
Document Loader
Handles loading and processing of multiple document types

This is like a smart file reader that can:
1. Read TXT files
2. Read PDF files (extracts text from pages)
3. Read DOCX files (extracts text from Word documents)
4. Keep track of metadata (filename, size, when modified)
"""
import sys
from pathlib import Path
from typing import List, Dict
from datetime import datetime

# Add project root to path so we can import from shared/
project_root = Path(__file__).parent.parent.parent.parent
sys.path.insert(0, str(project_root))

from shared.utils.logger import logger


class DocumentLoader:
    """
    Loads and processes documents of various types
    
    Think of this as your document reading assistant. You give it
    a folder of files, and it can read them all and tell you about them.
    """
    
    # These are the file types we can handle
    SUPPORTED_TYPES = ['.txt', '.md', '.pdf', '.docx']
    
    def __init__(self, documents_dir: str = None):
        """
        Initialize document loader
        
        Args:
            documents_dir: Path to folder containing documents
                          If None, uses the default data/documents folder
        """
        if documents_dir is None:
            # Default to our data/documents folder
            self.documents_dir = Path(__file__).parent.parent.parent / "data" / "documents"
        else:
            self.documents_dir = Path(documents_dir)
        
        logger.info(f"DocumentLoader initialized with directory: {self.documents_dir}")
    
    def load_text_file(self, file_path: Path) -> Dict:
        """
        Load a text file (TXT or MD)
        
        This is the simplest case - just read the text!
        
        Args:
            file_path: Path to the text file
            
        Returns:
            Dictionary with:
            - content: The actual text from the file
            - metadata: Info about the file (name, size, etc.)
        """
        try:
            # Open and read the file
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Collect information about this file
            metadata = {
                'source': str(file_path),           # Full path
                'filename': file_path.name,          # Just the filename
                'type': 'text',                      # File type
                'extension': file_path.suffix,       # .txt or .md
                'size_bytes': file_path.stat().st_size,  # How big is it?
                'modified': datetime.fromtimestamp(
                    file_path.stat().st_mtime
                ).isoformat()  # When was it last changed?
            }
            
            logger.info(f"Loaded text file: {file_path.name} ({len(content)} characters)")
            
            return {
                'content': content,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Error loading text file {file_path}: {str(e)}")
            raise
    
    def load_pdf_file(self, file_path: Path) -> Dict:
        """
        Load a PDF file
        
        PDFs are trickier - we need to extract text from each page.
        We use the pypdf library for this.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Dictionary with content and metadata
        """
        try:
            # Import PDF library (only when needed)
            from pypdf import PdfReader
            
            # Open the PDF
            reader = PdfReader(file_path)
            
            # Extract text from each page
            content_parts = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                # Add page number so we can cite sources
                content_parts.append(f"[Page {i+1}]\n{text}")
            
            # Join all pages together
            content = "\n\n".join(content_parts)
            
            # Metadata for PDFs includes page count
            metadata = {
                'source': str(file_path),
                'filename': file_path.name,
                'type': 'pdf',
                'extension': '.pdf',
                'pages': len(reader.pages),  # How many pages?
                'size_bytes': file_path.stat().st_size,
                'modified': datetime.fromtimestamp(
                    file_path.stat().st_mtime
                ).isoformat()
            }
            
            logger.info(f"Loaded PDF: {file_path.name} ({len(reader.pages)} pages)")
            
            return {
                'content': content,
                'metadata': metadata
            }
            
        except ImportError:
            logger.error("pypdf not installed. Install with: pip install pypdf")
            raise
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {str(e)}")
            raise
    
    def load_docx_file(self, file_path: Path) -> Dict:
        """
        Load a Word document (DOCX)
        
        Uses python-docx to extract text from Word documents.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary with content and metadata
        """
        try:
            # Import Word document library
            from docx import Document
            
            # Open the Word document
            doc = Document(file_path)
            
            # Extract text from each paragraph
            content_parts = [
                paragraph.text 
                for paragraph in doc.paragraphs 
                if paragraph.text  # Skip empty paragraphs
            ]
            content = "\n\n".join(content_parts)
            
            metadata = {
                'source': str(file_path),
                'filename': file_path.name,
                'type': 'docx',
                'extension': '.docx',
                'paragraphs': len(doc.paragraphs),
                'size_bytes': file_path.stat().st_size,
                'modified': datetime.fromtimestamp(
                    file_path.stat().st_mtime
                ).isoformat()
            }
            
            logger.info(f"Loaded DOCX: {file_path.name} ({len(doc.paragraphs)} paragraphs)")
            
            return {
                'content': content,
                'metadata': metadata
            }
            
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {str(e)}")
            raise
    
    def load_document(self, file_path: Path) -> Dict:
        """
        Load any supported document type
        
        This is the "smart router" - it looks at the file extension
        and calls the right loader function.
        
        Args:
            file_path: Path to any document
            
        Returns:
            Dictionary with content and metadata
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()  # .txt, .pdf, etc.
        
        # Check if we support this file type
        if extension not in self.SUPPORTED_TYPES:
            raise ValueError(
                f"Unsupported file type: {extension}. "
                f"Supported: {self.SUPPORTED_TYPES}"
            )
        
        # Route to the appropriate loader
        if extension in ['.txt', '.md']:
            return self.load_text_file(file_path)
        elif extension == '.pdf':
            return self.load_pdf_file(file_path)
        elif extension == '.docx':
            return self.load_docx_file(file_path)
    
    def load_all_documents(self) -> List[Dict]:
        """
        Load ALL documents from the documents directory
        
        This is super convenient - just load everything at once!
        
        Returns:
            List of document dictionaries
        """
        if not self.documents_dir.exists():
            logger.warning(f"Documents directory not found: {self.documents_dir}")
            return []
        
        documents = []
        
        # Loop through all files in the directory
        for file_path in self.documents_dir.iterdir():
            # Only process files (not folders) with supported extensions
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_TYPES:
                try:
                    doc = self.load_document(file_path)
                    documents.append(doc)
                except Exception as e:
                    logger.error(f"Failed to load {file_path.name}: {str(e)}")
                    continue  # Skip this file and continue with others
        
        logger.info(f"Loaded {len(documents)} documents from {self.documents_dir}")
        return documents
    
    def get_document_summary(self, documents: List[Dict]) -> Dict:
        """
        Get summary statistics for loaded documents
        
        Useful for understanding what you've loaded:
        - How many documents?
        - What types?
        - Total size?
        
        Args:
            documents: List of document dictionaries
            
        Returns:
            Summary statistics dictionary
        """
        if not documents:
            return {
                'total_documents': 0,
                'by_type': {},
                'total_size_bytes': 0
            }
        
        # Count documents by type
        by_type = {}
        total_size = 0
        
        for doc in documents:
            doc_type = doc['metadata']['type']
            by_type[doc_type] = by_type.get(doc_type, 0) + 1
            total_size += doc['metadata']['size_bytes']
        
        return {
            'total_documents': len(documents),
            'by_type': by_type,  # e.g., {'text': 3, 'pdf': 1}
            'total_size_bytes': total_size,
            'total_size_mb': round(total_size / (1024 * 1024), 2)
        }